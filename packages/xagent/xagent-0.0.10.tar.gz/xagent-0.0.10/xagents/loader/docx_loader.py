#!/usr/bin/env python
# -*- encoding: utf-8 -*-
'''
@Time    :   2024/04/12 15:00:12
@Author  :   ChenHao
@Description  :  doc/docx文档解析
@Contact :   jerrychen1990@gmail.com
'''



from typing import Iterable

import docx.table


from xagents.loader.utils import upload_to_minio, image2text
from xagents.config import *
from xagents.loader.common import Chunk, AbstractLoader, ImageChunk, TextChunk, TableChunk, get_image_name_path, merge_chunks

import docx
import pandas as pd

from PIL import Image
from docx.text.paragraph import Paragraph
import xml.etree.ElementTree as ET
from docx.document import Document as DOC
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell
from xml.etree import ElementTree
from io import StringIO
import io
import csv

def iter_block_items(parent):
    """
    Yield each paragraph and table child within *parent*, in document order.
    Each returned value is an instance of either Table or Paragraph. *parent*
    would most commonly be a reference to a main Document object, but
    also works for a _Cell object, which itself can contain paragraphs and tables.
    """
    if isinstance(parent, DOC):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield docx.table.Table(child, parent)

    
#This function extracts the table from the document object as a dataframe
def read_docx_tables(document, tab_id=None, **kwargs):
    """
    parse table(s) from a Word Document (.docx) into Pandas DataFrame(s)

    Parameters:
        filename:   file name of a Word Document

        tab_id:     parse a single table with the index: [tab_id] (counting from 0).
                    When [None] - return a list of DataFrames (parse all tables)

        kwargs:     arguments to pass to `pd.read_csv()` function

    Return: a single DataFrame if tab_id != None or a list of DataFrames otherwise
    """
    def read_docx_tab(tab, **kwargs):
        vf = io.StringIO()
        writer = csv.writer(vf)
        for row in tab.rows:
            writer.writerow(cell.text for cell in row.cells)
        vf.seek(0)
        return pd.read_csv(vf, **kwargs)

#    doc = Document(filename)
    if tab_id is None:
        return [read_docx_tab(tab, **kwargs) for tab in document.tables]
    else:
        try:
            return read_docx_tab(document.tables[tab_id], **kwargs)
        except IndexError:
            print('Error: specified [tab_id]: {}  does not exist.'.format(tab_id))
            raise

def analyse_page(document:DOC, file_name:str, upload_image:bool, ocr:bool)->Iterable[Chunk]:

    page_index = 1
    table_idx=0
    image_idx = 0

    for block in iter_block_items(document):
        if 'text' in str(block):
        
            content = str(block.text)
            yield TextChunk(content=content, page_idx=page_index)

            for run in block.runs:
                if 'pageBreakBefore' in run.element.xml or '\f' in run.text:
                    page_index += 1
                xmlstr = str(run.element.xml)
                my_namespaces = dict([node for _, node in ElementTree.iterparse(StringIO(xmlstr), events=['start-ns'])])
                root = ET.fromstring(xmlstr) 
                #Check if pic is there in the xml of the element. If yes, then extract the image data
                if 'pic:pic' in xmlstr:
                    for pic in root.findall('.//pic:pic', my_namespaces):
                        blip_elem = pic.find("pic:blipFill/a:blip", my_namespaces)
                        embed_attr = blip_elem.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                        document_part = document.part
                        image_part = document_part.related_parts[embed_attr]
                        image_name, image_path =get_image_name_path(file_name, 0, image_idx)    

                        if upload_image:
                            with open(image_path, "wb") as f:
                                f.write(image_part._blob)
                
                            url = upload_to_minio(image_path)
                        else:
                            url=None
                        content = image2text(Image.open(io.BytesIO(image_part._blob))) if ocr else None
                        # image_base64 = base64.b64encode(image_part._blob)
                        # image_base64 = image_base64.decode()        
                        yield ImageChunk(url=url, content=content,page_idx=page_index, image_name=image_name)
                    image_idx = image_idx + 1
                
        elif 'table' in str(block):
            content = str(block)
            df:pd.DataFrame = read_docx_tables(document=document,tab_id=table_idx)
            yield TableChunk(page_idx=page_index, data=df)
            table_idx+=1

            
class DOCXLoader(AbstractLoader):
    def __init__(self, max_page:int=None, **kwargs):
        """构建pdf加载器

        Args:
            max_page (int, optional): 最大页数. Defaults to None：不限定页数
            extract_images (bool, optional): 是否使用ocr抽取其中的图片. Defaults to False.
        """
        super().__init__(**kwargs)
        self.max_page = max_page


    def load(self, file_path: str, start_page=1, end_page=None, upload_image=True, ocr=False, **kwargs) -> Iterable[Chunk]:
        #TODO 缺少页码控制
        doc = docx.Document(file_path)
        chunks = analyse_page(doc, file_name=os.path.basename(file_path), upload_image=upload_image, ocr=ocr)
        chunks = merge_chunks(chunks)
        yield from chunks


if __name__ == "__main__":
    doc = docx.Document("data/kb_file/Xagents中间件.docx")
    chunks = analyse_page(doc)
    for chunk in chunks:
        print(chunk)