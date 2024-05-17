import pandas as pd
import altair as alt
from docx import Document
from fpdf import FPDF

class Report:
    """
    A class for generating reports containing data tables, charts, and qualitative analysis.

    Attributes:
    data (DataFrame): The input data for the report.
    """

    def __init__(self, data):
        """
        Initializes the Report class with input data.

        Args:
        data (DataFrame): The input DataFrame for the report.
        """
        self.data = data

    def generate_report(self, report_format='docx'):
        """
        Generates the report based on the specified format.

        Args:
        report_format (str, optional): The format of the report ('docx', 'pdf', 'latex', 'html'). Defaults to 'docx'.

        Returns:
        str: The report content.
        """
        if report_format == 'docx':
            return self._generate_docx_report()
        elif report_format == 'pdf':
            return self._generate_pdf_report()
        elif report_format == 'latex':
            return self._generate_latex_report()
        elif report_format == 'html':
            return self._generate_html_report()
        else:
            raise ValueError("Invalid report format. Supported formats: 'docx', 'pdf', 'latex', 'html'.")

    def _generate_docx_report(self):
        """
        Generates a DOCX report containing data tables, charts, and qualitative analysis.

        Returns:
        str: The DOCX report content.
        """
        document = Document()

        # Add data table
        table_name = "Data Table"
        self._add_table_to_docx(document, table_name, self.data)

        # Add qualitative analysis
        qualitative_analysis = self._generate_qualitative_analysis()
        document.add_paragraph("Qualitative Analysis:")
        document.add_paragraph(qualitative_analysis)

        # Save to DOCX file
        filename = "report.docx"
        document.save(filename)
        return filename

    def _add_table_to_docx(self, document, table_name, table_data):
        """
        Adds a table to a DOCX document.

        Args:
        document: The DOCX document object.
        table_name (str): The name of the table.
        table_data (DataFrame): The data to be added to the table.
        """
        document.add_paragraph(table_name)
        table = document.add_table(rows=1, cols=len(table_data.columns))
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(table_data.columns):
            hdr_cells[i].text = col
        for index, row in table_data.iterrows():
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                row_cells[i].text = str(val)

    def _generate_pdf_report(self):
        """
        Generates a PDF report containing data tables, charts, and qualitative analysis.

        Returns:
        str: The PDF report content.
        """
        pdf = FPDF()
        pdf.add_page()

        # Add data table
        table_name = "Data Table"
        self._add_table_to_pdf(pdf, table_name, self.data)

        # Add qualitative analysis
        qualitative_analysis = self._generate_qualitative_analysis()
        pdf.cell(200, 10, txt="Qualitative Analysis:", ln=True)
        pdf.multi_cell(0, 10, qualitative_analysis)

        # Save to PDF file
        filename = "report.pdf"
        pdf.output(filename)
        return filename

    def _add_table_to_pdf(self, pdf, table_name, table_data):
        """
        Adds a table to a PDF document.

        Args:
        pdf: The PDF document object.
        table_name (str): The name of the table.
        table_data (DataFrame): The data to be added to the table.
        """
        pdf.set_font("Arial", size=12)
        pdf.cell(200, 10, txt=table_name, ln=True)
        for index, row in table_data.iterrows():
            pdf.cell(200, 10, txt=row.to_string(), ln=True)

    def _generate_qualitative_analysis(self):
        """
        Generates qualitative analysis based on the input data.

        Returns:
        str: The qualitative analysis.
        """
        # Assuming you have implemented qualitative analysis generation
        analysis = "Qualitative analysis generated here."
        return analysis


