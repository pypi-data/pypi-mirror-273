import csv
from abc import abstractmethod
from copy import deepcopy
from pathlib import Path

from docx import Document
from openpyxl import Workbook

from .subtitle import Subtitle, SubtitleProcessor


class AbstractSubtitleSaver:
    def __init__(self, filename, keep_time=False):
        self.target = Path(filename)
        self.keep_time = keep_time
        self.file = None
        self.encoding = 'utf-8'
        self.processors: [SubtitleProcessor] = []

    def __enter__(self):
        self.file = open(self.target, 'w', encoding=self.encoding)
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        self.file.flush()
        self.file.close()
        return False

    def install_processor(self, processor):
        if isinstance(processor, SubtitleProcessor) and (processor not in self.processors):
            self.processors.append(processor)
        return self

    def process_subtitle(self, subtitle: Subtitle) -> Subtitle:
        result = deepcopy(subtitle)
        for p in self.processors:
            result = p(result)
        return result

    @abstractmethod
    def write(self, subtitle: Subtitle):
        pass


class TxtSaver(AbstractSubtitleSaver):
    LINE_TEMPLATE = {
        True: '{start}\t{content}\n',
        False: '{content}\n'
    }

    def __init__(self, target, keep_time=False):
        super(TxtSaver, self).__init__(target, keep_time)

    def write(self, subtitle: Subtitle):
        sub = self.process_subtitle(subtitle)
        line = TxtSaver.LINE_TEMPLATE[self.keep_time].format(start=sub.start.timestamp, content=sub.content)
        self.file.write(line)


class SrtSaver(AbstractSubtitleSaver):
    LINE_TEMPLATE = '{number}\n{start} --> {end}\n{content}\n\n'

    def __init__(self, target, keep_time=True):
        super(SrtSaver, self).__init__(target, keep_time)
        self.__count = 1

    def write(self, subtitle: Subtitle):
        sub = self.process_subtitle(subtitle)
        line = SrtSaver.LINE_TEMPLATE.format(
            number=self.__count,
            start=sub.start.timestamp,
            end=sub.end.timestamp,
            content=sub.content
        )
        self.file.write(line)
        self.__count += 1


class WordSaver(AbstractSubtitleSaver):
    LINE_TEMPLATE = {
        True: '{start}\t{content}',
        False: '{content}'
    }

    def __init__(self, target, keep_time=False):
        super(WordSaver, self).__init__(target, keep_time)
        self.document = Document()

    def __enter__(self):
        self.document.add_heading(Path(self.target.stem).name, level=1)
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        self.document.save(self.target)
        return False

    def write(self, subtitle: Subtitle):
        sub = self.process_subtitle(subtitle)
        line = self.LINE_TEMPLATE[self.keep_time].format(
            start=sub.start.timestamp,
            content=sub.content
        )
        self.document.add_paragraph(line)


class ExcelSaver(AbstractSubtitleSaver):
    """如果 keep_time 为 True，则会追加秒数队列"""
    HEADERS = {
        True: ['开始', '结束', '开始（秒）', '结束（秒）', '台词'],
        False: ['开始', '结束', '台词']
    }

    def __init__(self, target, keep_time=True):
        super(ExcelSaver, self).__init__(target, keep_time)
        self.workbook = Workbook()
        self.table = self.workbook.active

    def __enter__(self):
        self.table.title = '字幕表格'
        self.table.append(self.HEADERS[self.keep_time])
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        self.workbook.save(self.target)
        return False

    def write(self, subtitle: Subtitle):
        sub = self.process_subtitle(subtitle)
        row = [str(sub.start.timestamp), str(sub.end.timestamp)]
        if self.keep_time:
            row.append(sub.start.second)
            row.append(sub.end.second)
        row.append(sub.content)
        self.table.append(row)


class CsvSaver(AbstractSubtitleSaver):
    """如果 keep_time 为 True，则会追加秒数队列"""
    HEADERS = {
        True: ['开始', '结束', '开始（秒）', '结束（秒）', '台词'],
        False: ['开始', '结束', '台词']
    }

    def __init__(self, target, keep_time=False):
        super(CsvSaver, self).__init__(target, keep_time)
        self.__saver = None

    def __enter__(self):
        self.file = open(self.target, 'w', newline='', encoding=self.encoding)
        self.__saver = csv.DictWriter(self.file, self.HEADERS[self.keep_time], extrasaction='ignore')
        self.__saver.writeheader()
        return self

    def write(self, subtitle: Subtitle):
        sub = self.process_subtitle(subtitle)
        start, end = sub.start, sub.end
        row = {
            '开始': str(start.timestamp),
            '结束': str(end.timestamp),
            '开始（秒）': start.second,
            '结束（秒）': end.second,
            '台词': sub.content.strip()
        }
        self.__saver.writerow(row)
