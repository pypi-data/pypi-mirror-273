import csv
from abc import abstractmethod
from pathlib import Path

from docx import Document
from openpyxl import Workbook

from .subtitle import Subtitle


class AbstractSubtitleSaver:
    def __init__(self, filename, keep_time=False):
        self.target = Path(filename)
        self.keep_time = keep_time
        self.file = None
        self.encoding = 'utf-8'

    def __enter__(self):
        self.file = open(self.target, 'w', encoding=self.encoding)
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        self.file.flush()
        self.file.close()
        return False

    @abstractmethod
    def write(self, subtitle: Subtitle):
        pass


class TxtSaver(AbstractSubtitleSaver):
    LINE_TEMPLATE = {
        True: '{start}\t{content}\n',
        False: '{content}\n'
    }

    def __init__(self, target, keep_time=False):
        super(TxtSaver, self).__init__(target, keep_time)

    def write(self, subtitle: Subtitle):
        line = TxtSaver.LINE_TEMPLATE[self.keep_time].format(start=subtitle.start.timestamp, content=subtitle.content)
        self.file.write(line)


class SrtSaver(AbstractSubtitleSaver):
    LINE_TEMPLATE = '{number}\n{start} --> {end}\n{content}\n\n'

    def __init__(self, target, keep_time=True):
        super(SrtSaver, self).__init__(target, keep_time)
        self.__count = 1

    def write(self, subtitle: Subtitle):
        line = SrtSaver.LINE_TEMPLATE.format(
            number=self.__count,
            start=subtitle.start.timestamp,
            end=subtitle.end.timestamp,
            content=subtitle.content
        )
        self.file.write(line)
        self.__count += 1


class WordSaver(AbstractSubtitleSaver):
    LINE_TEMPLATE = {
        True: '{start}\t{content}',
        False: '{content}'
    }

    def __init__(self, target, keep_time=False):
        super(WordSaver, self).__init__(target, keep_time)
        self.document = Document()

    def __enter__(self):
        self.document.add_heading(Path(self.target.stem).name, level=1)
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        self.document.save(self.target)
        return False

    def write(self, subtitle: Subtitle):
        line = self.LINE_TEMPLATE[self.keep_time].format(
            start=subtitle.start.timestamp,
            content=subtitle.content
        )
        self.document.add_paragraph(line)


class ExcelSaver(AbstractSubtitleSaver):
    """如果 keep_time 为 True，则会追加秒数队列"""
    HEADERS = {
        True: ['开始', '结束', '开始（秒）', '结束（秒）', '台词'],
        False: ['开始', '结束', '台词']
    }

    def __init__(self, target, keep_time=True):
        super(ExcelSaver, self).__init__(target, keep_time)
        self.workbook = Workbook()
        self.table = self.workbook.active

    def __enter__(self):
        self.table.title = '字幕表格'
        self.table.append(self.HEADERS[self.keep_time])
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        self.workbook.save(self.target)
        return False

    def write(self, subtitle: Subtitle):
        row = [str(subtitle.start.timestamp), str(subtitle.end.timestamp)]
        if self.keep_time:
            row.append(subtitle.start.second)
            row.append(subtitle.end.second)
        row.append(subtitle.content)
        self.table.append(row)


class CsvSaver(AbstractSubtitleSaver):
    """如果 keep_time 为 True，则会追加秒数队列"""
    HEADERS = {
        True: ['开始', '结束', '开始（秒）', '结束（秒）', '台词'],
        False: ['开始', '结束', '台词']
    }

    def __init__(self, target, keep_time=False):
        super(CsvSaver, self).__init__(target, keep_time)
        self.__saver = None

    def __enter__(self):
        self.file = open(self.target, 'w', newline='', encoding=self.encoding)
        self.__saver = csv.DictWriter(self.file, self.HEADERS[self.keep_time], extrasaction='ignore')
        self.__saver.writeheader()
        return self

    def write(self, subtitle: Subtitle):
        start, end = subtitle.start, subtitle.end
        row = {
            '开始': str(start.timestamp),
            '结束': str(end.timestamp),
            '开始（秒）': start.second,
            '结束（秒）': end.second,
            '台词': subtitle.content.strip()
        }
        self.__saver.writerow(row)
