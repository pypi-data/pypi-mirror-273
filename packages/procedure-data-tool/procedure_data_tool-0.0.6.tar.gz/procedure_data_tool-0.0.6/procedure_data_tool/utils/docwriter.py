from docx import Document
from collections import OrderedDict
from docx.shared import Pt
from docx.shared import RGBColor
from procedure_data_tool.utils.valve2 import Valve2
from procedure_data_tool.utils.valve3 import Valve3

class DocWriter():
    def __init__(self, name="MyDoc"):
        self.name = name
        self.doc = Document()
        run = self.doc.add_heading("", 1).add_run()
        font = run.font
        font.name = "Times New Roman"
        font.size = Pt(11)
        font.color.rgb = RGBColor(0x42,0x42,0x42)
        run.add_text(self.name)
        
    def makeSection(self, name, instruction = None):
        paragraph = self.doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph.paragraph_format.space_before = Pt(1)
        section_name = paragraph.add_run(name)
        section_name.font.bold = True
        section_name.font.name = "Times New Roman"
        section_name.font.size = Pt(11)
        prompt = paragraph.add_run(instruction)
        prompt.font.name = "Times New Roman"
        prompt.font.size = Pt(11)
        return paragraph

    def save(self, filename= "PrintedRoute.docx"):
        self.doc.save(filename)

    def buildDocument(self, route, pits):
        route_list = self.makeSection("Valves in Route (reference only): ", "DVI Credited YES/NO/Position dependent")
        used_pits = OrderedDict()
        used_jumpers = OrderedDict()
        for node in route:
            if node.pit in pits:
                used_pits[node.pit] = None
            if node.onJumper:
                jumper = (node.pit, node.jumper)
                used_jumpers[jumper] = None
            if node.show:
                route_list.add_run("\n")
                route_list.add_run(node.EIN())
                route_list.add_run("\t")
                route_list.add_run(node.dvi_credited)
        
        heaterEINs = self.makeSection("Section 5.5.3 heaters: " ,"Replace existing data with the following:")
        for pit in used_pits:
            for heater in pits[pit].heaters:
                heaterEINs.add_run("\n")      
                heaterEINs.add_run(heater)
                heaterEINs.add_run("\t \t")
                heaterEINs.add_run(pits[pit].nacePMID)
        
        pits579 = self.makeSection("Steps 5.17.9: ","Replace existing data with the following:")
        for pit in used_pits:
            pits579.add_run("\n")
            pits579.add_run(pits[pit].label)
        checklist1 = self.makeSection("Checklist 1: ","Replace list with:")
        for jumper in used_jumpers:
            checklist1.add_run("\n")
            checklist1.add_run(jumper[0])
            checklist1.add_run("\t \t")
            checklist1.add_run("Jumper: ").font.bold = True
            checklist1.add_run(jumper[1])
        checklist3 = self.makeSection("Checklist 3:","")
        checklist4 = self.makeSection("Checklist 4:","")
        checklist5 = self.makeSection("Checklist 5:","")
        checklist6 = self.makeSection("Checklist 6:","")
        checklist7LD = self.makeSection("Checklist 7 - Tank pit/Structure Leak Detection")
        checklist7TF = self.makeSection("Checklist 7 - TFSPS Temperature Equipment Checks")
        for pit in used_pits:
            for tfsps , pmid in zip(pits[pit].tfsps,pits[pit].tfsps_pmid):
                checklist7TF.add_run("\n")
                checklist7TF.add_run(tfsps)
                checklist7TF.add_run("\t \t")
                checklist7TF.add_run(pmid)

        checklist7D = self.makeSection("Checklist 7 - Drain Seal Assemblies:")
        for pit in used_pits:
            checklist7D.add_run("\n")
            checklist7D.add_run(pits[pit].label)
            checklist7D.add_run(" ")
            checklist7D.add_run(pits[pit].drain)
            checklist7D.add_run("\t \t")
            checklist7D.add_run(pits[pit].drainSealPos)

        checklist7N = self.makeSection("Checklist 7 - NACE Inspection:")
        for pit in used_pits:
            checklist7N.add_run("\n")
            checklist7N.add_run(pits[pit].nace)
            checklist7N.add_run("\t \t")
            checklist7N.add_run(pits[pit].nacePMID)

        dviBlockTest = self.makeSection("Test: ", "(Valve Positions Test)")
        for i in range(1,len(route)-2):
            if (type(route[i]) == Valve3 or type(route[i]) == Valve2 ):
                dviBlockTest.add_run("\n")
                dviBlockTest.add_run(route[i].EIN())
                dviBlockTest.add_run("\t \t")
                dviBlockTest.add_run(route[i].position)


